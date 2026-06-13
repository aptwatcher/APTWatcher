"""
Bilingual ``.docx`` campaign report renderer.

Phase 3.8 task #63. Produces the operator-facing Word deliverable. Two
languages are supported (``en``, ``fr``); both renders share identical
section ordering and table shapes per the design doc. Only the label
strings differ, sourced from :mod:`core.analysis.i18n`.

Design rules:

- **Python-docx backed.** No ``.docx`` templates on disk; every style is
  programmatically created so the renderer has no hidden dependencies.
- **No overwrite.** ``output_path`` must not exist. The renderer raises
  :class:`ReportRenderError` rather than silently clobbering prior
  artifacts -- the offline-to-online handoff relies on ``.docx`` files
  being hashed once and never rewritten.
- **Severity from confidence.** The data model carries ``confidence``
  (0..1) on each ``Finding`` but no explicit severity. Severity bands
  are derived with a small monotonic rule so the table and the summary
  stay consistent.
- **No operator prose is translated.** The French render reuses
  English-language ``Finding.summary`` and IOC values verbatim -- only
  headings, labels, and boilerplate are swapped. Missing translations
  for operator prose are an explicit design choice (see the design
  doc, "Bilingual .docx" section).

Reference:

- ``docs/design/analysis-output-pipeline.md`` -- "Report renderers"
  -> "Campaign .docx (EN/FR)".
- ``core.analysis.i18n`` -- translation table.
"""

from __future__ import annotations

from datetime import UTC, datetime
from pathlib import Path
from typing import Any

from core.analysis.i18n import Language, translate
from core.types import Finding, FindingCitation, IOCVerdict, utcnow


class ReportRenderError(ValueError):
    """Raised when the finding/IOC set cannot produce a valid report."""


# -----------------------------------------------------------------------
# Severity mapping
# -----------------------------------------------------------------------


def _severity_band(confidence: float) -> str:
    """
    Map a ``Finding.confidence`` float to a named severity band.

    Bands are intentionally coarse (four levels plus ``none``) so the
    rendered table reads at a glance. The thresholds are chosen so the
    most dangerous findings (confidence >= 0.9) are unmistakably red.
    """
    if confidence >= 0.9:
        return "critical"
    if confidence >= 0.75:
        return "high"
    if confidence >= 0.5:
        return "medium"
    if confidence > 0.0:
        return "low"
    return "none"


_SEVERITY_ORDER: dict[str, int] = {
    "critical": 4,
    "high": 3,
    "medium": 2,
    "low": 1,
    "none": 0,
}


def _localized_severity(band: str, language: Language) -> str:
    """Return the translated severity label for the given band name."""
    return translate(f"severity_{band}", language=language)


# -----------------------------------------------------------------------
# Helpers
# -----------------------------------------------------------------------


def _format_datetime(dt: datetime) -> str:
    """Format a datetime as an ISO-8601 UTC string with a trailing ``Z``."""
    if dt.tzinfo is None:
        dt = dt.replace(tzinfo=UTC)
    dt = dt.astimezone(UTC)
    return dt.strftime("%Y-%m-%dT%H:%M:%SZ")


def _highest_severity(findings: list[Finding]) -> str:
    """Return the name of the highest severity band across ``findings``."""
    if not findings:
        return "none"
    best = "none"
    for f in findings:
        band = _severity_band(f.confidence)
        if _SEVERITY_ORDER[band] > _SEVERITY_ORDER[best]:
            best = band
    return best


def _count_iocs_by_type(iocs: list[IOCVerdict]) -> dict[str, int]:
    """Return a per-type tally of IOCs, sorted by type name for stability."""
    counts: dict[str, int] = {}
    for ioc in iocs:
        counts[ioc.ioc_type] = counts.get(ioc.ioc_type, 0) + 1
    return dict(sorted(counts.items()))


def _citation_summary(citation: FindingCitation, language: Language) -> str:
    """Render a single citation as a compact inline string."""
    parts: list[str] = [f"{translate('label_citation_source', language=language)}: {citation.source}"]
    if citation.locator:
        parts.append(f"{translate('label_citation_locator', language=language)}: {citation.locator}")
    if citation.tool_call_id:
        parts.append(
            f"{translate('label_citation_tool_call', language=language)}: {citation.tool_call_id}",
        )
    return " | ".join(parts)


# -----------------------------------------------------------------------
# Public API
# -----------------------------------------------------------------------


def render_docx_report(
    *,
    findings: list[Finding],
    iocs: list[IOCVerdict],
    output_path: Path,
    incident_id: str,
    campaign_tag: str,
    language: Language = "en",
    operator: str | None = None,
    generated_at: datetime | None = None,
) -> Path:
    """
    Render a professional incident campaign report as a ``.docx`` file.

    Parameters
    ----------
    findings:
        Verified findings for the incident. May be empty; in that case
        the executive summary carries a "no findings" placeholder and
        the findings table is omitted.
    iocs:
        Indicators of compromise associated with the incident.
    output_path:
        Filesystem path to write. Must NOT exist -- the renderer refuses
        to overwrite.
    incident_id:
        Run identifier (matches the audit log).
    campaign_tag:
        Human-readable campaign slug (e.g. ``S01_RANSOMWARE_TRIAGE``).
    language:
        ``"en"`` (default) or ``"fr"``.
    operator:
        Optional operator identifier. Shown on the title page. When
        ``None`` a localized "unknown" placeholder is used.
    generated_at:
        Optional override for the "generated" timestamp (used by tests
        for determinism). Defaults to :func:`core.types.utcnow`.

    Returns
    -------
    Path
        The same ``output_path`` value, for call-site chaining.
    """
    # Defer the import so the analysis package stays importable on hosts
    # where python-docx is not installed (e.g. lightweight offline shells
    # that only consume rule generators). The failure is surfaced here
    # as a clean ``ReportRenderError``.
    try:
        from docx import Document
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.shared import Pt
    except ImportError as exc:
        msg = "python-docx is required to render .docx reports; install with `pip install python-docx`"
        raise ReportRenderError(msg) from exc

    if language not in ("en", "fr"):
        msg = f"unsupported language: {language!r}"
        raise ReportRenderError(msg)

    if not incident_id:
        raise ReportRenderError("incident_id must be a non-empty string")
    if not campaign_tag:
        raise ReportRenderError("campaign_tag must be a non-empty string")

    output_path = Path(output_path)
    if output_path.exists():
        msg = f"refusing to overwrite existing report: {output_path}"
        raise ReportRenderError(msg)
    output_path.parent.mkdir(parents=True, exist_ok=True)

    when = generated_at or utcnow()
    operator_label = operator or translate("label_unknown_operator", language=language)

    doc = Document()

    # Shrink default body font slightly for table density.
    style = doc.styles["Normal"]
    style.font.size = Pt(11)

    # ---- Title page -------------------------------------------------
    title_para = doc.add_heading(translate("title_campaign_report", language=language), level=0)
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    meta_table = doc.add_table(rows=4, cols=2)
    meta_table.style = "Light Grid Accent 1"
    _set_row(meta_table, 0, translate("label_campaign", language=language), campaign_tag)
    _set_row(meta_table, 1, translate("label_incident", language=language), incident_id)
    _set_row(meta_table, 2, translate("label_generated_at", language=language), _format_datetime(when))
    _set_row(meta_table, 3, translate("label_operator", language=language), operator_label)

    doc.add_paragraph()  # spacer

    # ---- Executive summary ------------------------------------------
    doc.add_heading(translate("heading_executive_summary", language=language), level=1)

    if findings:
        ioc_counts = _count_iocs_by_type(iocs)
        ioc_breakdown = ", ".join(f"{t}: {n}" for t, n in ioc_counts.items()) or translate(
            "placeholder_none", language=language,
        )
        highest_band = _highest_severity(findings)
        summary_text = (
            f"{translate('summary_prefix', language=language)} {len(findings)} "
            f"{translate('summary_finding_count', language=language)}"
            f" / {len(iocs)} {translate('summary_ioc_count', language=language)} "
            f"({ioc_breakdown}). "
            f"{translate('summary_highest_severity', language=language)}: "
            f"{_localized_severity(highest_band, language=language)}."
        )
        doc.add_paragraph(summary_text)
    else:
        doc.add_paragraph(translate("summary_no_findings", language=language))

    # ---- Findings table ---------------------------------------------
    doc.add_heading(translate("heading_findings", language=language), level=1)
    if findings:
        table = doc.add_table(rows=1, cols=5)
        table.style = "Light Grid Accent 1"
        headers = [
            translate("col_finding_id", language=language),
            translate("col_finding_title", language=language),
            translate("col_finding_severity", language=language),
            translate("col_finding_mitre", language=language),
            translate("col_finding_citations", language=language),
        ]
        for idx, text in enumerate(headers):
            table.rows[0].cells[idx].text = text
        for f in findings:
            row = table.add_row().cells
            row[0].text = f.finding_id
            row[1].text = f.summary
            row[2].text = _localized_severity(_severity_band(f.confidence), language=language)
            row[3].text = ", ".join(f.mitre) or translate("placeholder_none", language=language)
            row[4].text = str(len(f.evidence))
    else:
        doc.add_paragraph(translate("placeholder_none", language=language))

    # ---- IOC table ---------------------------------------------------
    doc.add_heading(translate("heading_iocs", language=language), level=1)
    if iocs:
        table = doc.add_table(rows=1, cols=4)
        table.style = "Light Grid Accent 1"
        headers = [
            translate("col_ioc_type", language=language),
            translate("col_ioc_value", language=language),
            translate("col_ioc_confidence", language=language),
            translate("col_ioc_source", language=language),
        ]
        for idx, text in enumerate(headers):
            table.rows[0].cells[idx].text = text
        for ioc in iocs:
            row = table.add_row().cells
            row[0].text = ioc.ioc_type
            row[1].text = ioc.value
            row[2].text = (
                f"{ioc.confidence:.2f}" if ioc.confidence is not None
                else translate("placeholder_n_a", language=language)
            )
            src_names = [s.name for s in ioc.sources] or [translate("placeholder_none", language=language)]
            row[3].text = ", ".join(src_names)
    else:
        doc.add_paragraph(translate("placeholder_none", language=language))

    # ---- Appendix: citation references ------------------------------
    doc.add_heading(translate("heading_appendix_citations", language=language), level=1)
    any_citation = False
    for f in findings:
        if not f.evidence:
            continue
        any_citation = True
        p = doc.add_paragraph()
        p.add_run(f"{f.finding_id}: ").bold = True
        doc.add_paragraph(
            "\n".join(_citation_summary(c, language=language) for c in f.evidence),
        )
    if not any_citation:
        doc.add_paragraph(translate("placeholder_none", language=language))

    doc.save(str(output_path))
    return output_path


# -----------------------------------------------------------------------
# Internal helpers that depend on python-docx types; kept at module level
# to keep ``render_docx_report`` readable.
# -----------------------------------------------------------------------


def _set_row(table: Any, row_idx: int, label: str, value: str) -> None:
    """Write a label/value pair into an existing two-column metadata row."""
    cells = table.rows[row_idx].cells
    cells[0].text = label
    cells[1].text = value


__all__ = [
    "ReportRenderError",
    "render_docx_report",
]
