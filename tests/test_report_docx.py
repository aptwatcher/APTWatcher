"""
Tests for the bilingual .docx campaign report renderer.

Coverage:
- Produces a real .docx file that re-opens via python-docx.
- EN and FR variants differ by heading text.
- Refuses to overwrite an existing file.
- Findings table row count equals len(findings) + 1 header row.
- IOC table row count equals len(iocs) + 1 header row.
- Empty findings list renders a "no findings" placeholder.
- Invalid language raises ReportRenderError.
- Metadata block carries campaign, incident, generated, operator.
"""

from __future__ import annotations

from datetime import UTC, datetime
from pathlib import Path

import pytest

from core.analysis.i18n import translate
from core.analysis.report_docx import ReportRenderError, render_docx_report
from core.types import Finding, FindingCitation, IOCProviderResult, IOCVerdict


def _ioc(value: str, ioc_type: str, confidence: float | None = 0.8) -> IOCVerdict:
    return IOCVerdict(
        value=value,
        ioc_type=ioc_type,  # type: ignore[arg-type]
        verdict="malicious",
        confidence=confidence,
        sources=[IOCProviderResult(name="stub", verdict="malicious", score=confidence)],
    )


def _findings() -> list[Finding]:
    return [
        Finding(
            finding_id="F-100",
            summary="Suspicious PowerShell execution",
            mitre=["T1059.001"],
            confidence=0.95,
            evidence=[FindingCitation(source="Security.evtx", locator="event_id=4104", tool_call_id="tc-1")],
        ),
        Finding(
            finding_id="F-101",
            summary="Scheduled task persistence",
            mitre=["T1053.005"],
            confidence=0.6,
            evidence=[FindingCitation(source="registry:HKLM\\...\\Schedule")],
        ),
    ]


def _iocs() -> list[IOCVerdict]:
    return [
        _ioc("203.0.113.10", "ipv4"),
        _ioc("evil.example", "domain"),
        _ioc("a" * 64, "sha256", confidence=0.55),
    ]


def _open_docx(path: Path):
    from docx import Document

    return Document(str(path))


def test_renders_valid_docx_en(tmp_path: Path) -> None:
    out = tmp_path / "Campaign_Report_TEST_20260419.docx"
    result = render_docx_report(
        findings=_findings(),
        iocs=_iocs(),
        output_path=out,
        incident_id="inc-001",
        campaign_tag="S01_TEST",
        language="en",
        operator="analyst",
        generated_at=datetime(2026, 4, 19, 15, 30, tzinfo=UTC),
    )
    assert result == out
    assert out.exists() and out.stat().st_size > 0

    doc = _open_docx(out)
    # The doc must parse and carry paragraphs.
    assert len(doc.paragraphs) > 0
    text_blob = "\n".join(p.text for p in doc.paragraphs)
    # English heading is present.
    assert translate("heading_executive_summary", language="en") in text_blob
    assert translate("heading_findings", language="en") in text_blob


def test_renders_valid_docx_fr(tmp_path: Path) -> None:
    out = tmp_path / "Campaign_Report_TEST_20260419_FR.docx"
    render_docx_report(
        findings=_findings(),
        iocs=_iocs(),
        output_path=out,
        incident_id="inc-001",
        campaign_tag="S01_TEST",
        language="fr",
    )
    doc = _open_docx(out)
    text_blob = "\n".join(p.text for p in doc.paragraphs)
    assert translate("heading_executive_summary", language="fr") in text_blob
    # English summary heading must NOT leak into the FR file.
    assert translate("heading_executive_summary", language="en") not in text_blob


def test_refuses_overwrite(tmp_path: Path) -> None:
    out = tmp_path / "report.docx"
    render_docx_report(
        findings=_findings(),
        iocs=_iocs(),
        output_path=out,
        incident_id="inc-001",
        campaign_tag="S01_TEST",
    )
    with pytest.raises(ReportRenderError, match="refusing to overwrite"):
        render_docx_report(
            findings=_findings(),
            iocs=_iocs(),
            output_path=out,
            incident_id="inc-001",
            campaign_tag="S01_TEST",
        )


def test_findings_table_row_count_matches_input(tmp_path: Path) -> None:
    out = tmp_path / "report.docx"
    findings = _findings()
    render_docx_report(
        findings=findings,
        iocs=_iocs(),
        output_path=out,
        incident_id="inc-001",
        campaign_tag="S01_TEST",
    )
    doc = _open_docx(out)
    # Meta table + findings table + iocs table expected.
    # Findings table is the first 5-column table.
    findings_tables = [t for t in doc.tables if len(t.columns) == 5]
    assert len(findings_tables) == 1
    ftable = findings_tables[0]
    # Header row + one row per finding.
    assert len(ftable.rows) == len(findings) + 1


def test_ioc_table_row_count_matches_input(tmp_path: Path) -> None:
    out = tmp_path / "report.docx"
    iocs = _iocs()
    render_docx_report(
        findings=_findings(),
        iocs=iocs,
        output_path=out,
        incident_id="inc-001",
        campaign_tag="S01_TEST",
    )
    doc = _open_docx(out)
    ioc_tables = [t for t in doc.tables if len(t.columns) == 4]
    assert len(ioc_tables) == 1
    itable = ioc_tables[0]
    assert len(itable.rows) == len(iocs) + 1


def test_empty_findings_renders_placeholder(tmp_path: Path) -> None:
    out = tmp_path / "empty.docx"
    render_docx_report(
        findings=[],
        iocs=[],
        output_path=out,
        incident_id="inc-empty",
        campaign_tag="EMPTY",
        language="en",
    )
    doc = _open_docx(out)
    text_blob = "\n".join(p.text for p in doc.paragraphs)
    # The "no findings" boilerplate must appear.
    assert translate("summary_no_findings", language="en") in text_blob


def test_invalid_language_raises(tmp_path: Path) -> None:
    out = tmp_path / "bad.docx"
    with pytest.raises(ReportRenderError, match="unsupported language"):
        render_docx_report(
            findings=_findings(),
            iocs=_iocs(),
            output_path=out,
            incident_id="inc-001",
            campaign_tag="S01_TEST",
            language="de",  # type: ignore[arg-type]
        )


def test_metadata_carries_ids_and_operator(tmp_path: Path) -> None:
    out = tmp_path / "meta.docx"
    render_docx_report(
        findings=_findings(),
        iocs=_iocs(),
        output_path=out,
        incident_id="inc-XYZ",
        campaign_tag="CAMPAIGN_META",
        operator="daniel",
        generated_at=datetime(2026, 1, 2, 3, 4, 5, tzinfo=UTC),
    )
    doc = _open_docx(out)
    # The metadata cover table is the first one with 2 columns.
    meta_table = next(t for t in doc.tables if len(t.columns) == 2)
    rows_text = [(r.cells[0].text, r.cells[1].text) for r in meta_table.rows]
    values = dict(rows_text)
    assert "inc-XYZ" in values.values()
    assert "CAMPAIGN_META" in values.values()
    assert "daniel" in values.values()
    # Timestamp rendered in the chosen format.
    assert "2026-01-02T03:04:05Z" in values.values()


def test_empty_incident_id_raises(tmp_path: Path) -> None:
    with pytest.raises(ReportRenderError, match="incident_id"):
        render_docx_report(
            findings=_findings(),
            iocs=_iocs(),
            output_path=tmp_path / "x.docx",
            incident_id="",
            campaign_tag="X",
        )
